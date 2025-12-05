import streamlit as st
from jinja2 import Environment, FileSystemLoader
import pdfkit
import os
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

import shutil

# Detect wkhtmltopdf availability
WKHTMLTOPDF_PATH = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
WKHTMLTOPDF_EXISTS = bool(shutil.which("wkhtmltopdf")) or os.path.exists(WKHTMLTOPDF_PATH)

if WKHTMLTOPDF_EXISTS:
    config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH) if os.path.exists(WKHTMLTOPDF_PATH) else pdfkit.configuration()
else:
    config = None




# Load templates from the "templates" folder
env = Environment(loader=FileSystemLoader("templates"))

# Explicit path to wkhtmltopdf.exe for Windows
path_wkhtmltopdf = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)

# Helper: DOCX export
def export_docx(name, email, phone, summary, skills, experience, education, output_path):
    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Key Skills", level=1)
    for skill in skills:
        doc.add_paragraph(skill, style="List Bullet")

    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph(experience)

    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)

    doc.save(output_path)

# Helper: Job description match score
def match_score(resume_text, job_text):
    vectorizer = TfidfVectorizer().fit([resume_text, job_text])
    vectors = vectorizer.transform([resume_text, job_text])
    score = cosine_similarity(vectors[0], vectors[1])[0][0]
    return round(score * 100, 2)

# Streamlit page setup
st.set_page_config(page_title="AI Resume Builder", layout="centered")
st.title("📄 AI Resume Builder")

# Input form
with st.form("resume_form"):
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    summary = st.text_area("Professional Summary")
    skills = st.text_area("Skills (comma-separated)")
    experience = st.text_area("Work Experience")
    education = st.text_area("Education")
    job_description = st.text_area("Paste Job Description (optional)")
    template_choice = st.selectbox(
        "Choose a Template",
        ["resume_template1.html", "resume_template2.html", "resume_template3.html"]
    )
    submitted = st.form_submit_button("Generate Resume")

if submitted:
    # Prepare data
    skills_list = [skill.strip() for skill in skills.split(",") if skill.strip()]
    template = env.get_template(template_choice)
    html = template.render(
        name=name,
        email=email,
        phone=phone,
        summary=summary,
        skills=skills_list,
        experience=experience,
        education=education
    )

    # Preview in Streamlit
    st.markdown("### 💡 Resume Preview")
    st.components.v1.html(html, height=800, scrolling=True)

    # Export options
    st.markdown("### 📥 Download Options")

    # HTML download
    st.download_button("Download HTML", html, file_name="resume.html", mime="text/html")

    # PDF export using wkhtmltopdf
   # pdf_path = "resume_styled.pdf"
    #pdfkit.from_string(html, pdf_path, configuration=config)
   # with open(pdf_path, "rb") as f:
   #     st.download_button("Download Styled PDF", f, file_name="resume.pdf", mime="application/pdf")
        
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
import os

# Set up Jinja2 environment to load templates
env = Environment(loader=FileSystemLoader("templates"))

def export_resume_to_pdf(template_name, context, output_path="outputs/resume.pdf"):
    """
    Render a Jinja2 HTML template with context data and export to PDF using WeasyPrint.
    
    Args:
        template_name (str): The name of the HTML template file (e.g., 'resume_template.html').
        context (dict): Dictionary of values to inject into the template (e.g., name, skills).
        output_path (str): Path where the PDF will be saved.
    """
    # Render HTML from template
    template = env.get_template(template_name)
    html_content = template.render(context)

    # Export to PDF
    HTML(string=html_content).write_pdf(output_path)

    print(f"✅ Resume exported successfully to {output_path}")

        
        

    # DOCX export
    docx_path = "resume.docx"
    export_docx(name, email, phone, summary, skills_list, experience, education, docx_path)
    with open(docx_path, "rb") as f:
        st.download_button("Download DOCX", f, file_name="resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # Job description match score
    if job_description.strip():
        resume_text = f"{summary} {', '.join(skills_list)} {experience} {education}"
        score = match_score(resume_text, job_description)
        st.markdown(f"### 📊 Job Match Score: {score}%")
